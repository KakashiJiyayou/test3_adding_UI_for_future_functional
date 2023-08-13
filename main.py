import os
import  sys

from module import package_management as _PM


import ntpath
import random

from PyQt5.QtWidgets import  *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from functools import partial
from resource_ui import Ui_MainWindow

import shutil
import time
import json
import datetime
import subprocess
import traceback, sys

from docx import Document
from PIL import Image
from docxcompose.composer import Composer


try:
    from ctypes import windll  # Only exists on Windows.
    myappid = 'gq.nigborobotic.filemanagment'
    windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
except ImportError:
    pass

try:
    from module import upload as M_upload
    from module import worker_pyqt as WorkQT
    from module import database as Database
    from module import runningdb as DatabaseRunning
    from module import database_management as DB_MGM
    from module import Create_Menu_From_MD as Read_Write_Menu
    from bypy import  ByPy

    bp = ByPy()
except:
    traceback.print_exc() 

"""
#NOTE Main Window---------------------------------Main Window------------------------------------------>  
"""
class MainWindow( QMainWindow ):

    ## Initial Function
    def __init__(self):
        super(MainWindow, self).__init__()
        self.threadpool = QThreadPool()
        self._IS_CREATING_MENU = False
        self.ui     =   Ui_MainWindow()

        self.ui.setupUi(self)
        # getting ui elements
        self.scroll_content_0 = self.ui.scrollArea_2_edit_page
        self.scroll_content_generate = self.ui.scrollArea_generate_page

        self.ui.icon_only_widget.hide()
        self.ui.home_btn_2.setChecked(1)

        # initialize variables
        # Variables with "_" prefix will be shared between pages
        self._current_page_index = 0
        
        self._folder_chosen = True
        self._is_english = True

        self.Upload_Button_Clicked = False
        self.Update_Button_Clicked = False

        self._search_path_list = []

       # Initialize an empty list of checkboxes
        self._checkboxes = []

        # Initialize new_item_list as an empty list
        self.new_item_list_generate_doc = []

        # initialize some module value
        self.M_Uplaod = M_upload

        # prevent creating multiple search reasult at the same time
        # for genrating checkbox
        self._IS_ADDING_CHECKBOXES = False

        # for left side generating doc list
        self.generate_doc_list =  []

        # generate user ID
        self._USER_ID = ''.join(random.choice('ABCDEFGHIJKLMNPQRSTUVWXYZ123456789') for _ in range(5))

        # just setting up the app
        self.setting_up_app()


    # NOTE - search abutton actions
    ## Function For Searching
    def on_search_btn_pressed(self):
        self.get_search_path_list_from_db ( None )

        if self._IS_NORMAL_PAGE:

            string_value = ""

            for item in self._search_path_list:
                string_value = string_value + str ( item ) + "\n"

            self.ui.plainText_show.setPlainText ( string_value  ) 
            # print ("normal page search button clicked")

        # show check box in 
        else :
            
            # put value in 
            self.new_item_list_generate_doc = self._search_path_list
            
            # print ( self.new_item_list_generate_doc )
            if not self._IS_ADDING_CHECKBOXES:
                self.update_list_for_generate_page ()
            # print ( "generate page search button")

    # NOTE - updating Checkboxes
    def update_list_for_generate_page(self):
        self._IS_ADDING_CHECKBOXES = True

        self.disable_g1 ()
        self.disable_ge2 ()
        self.diasble_side_bar ()

        

        try:  
            checkboxes_to_delete = []          
            print ("try clearing the layout parent _checkbox")
            # Clear existing checkboxes
            for checkbox in self._checkboxes:
                checkbox.setParent(None)
            self._checkboxes.clear()
        except Exception as e:
            print("Caught an exception: 2 ", e)
        
        self._checkboxes = []

        time.sleep (0.5)

        # Clear existing checkboxes
        try :
            print ("try clearing the layout parent widget scrollAreaWidgetContents_checkbox")
            layout = self.ui.scrollarea_checkbox_dirlist_000
            while layout.count():
                child = layout.takeAt(0)
                if child.widget():
                    child.widget().deleteLater()             
        except Exception as e:
            print("Caught an exception: 1", e)

        time.sleep(0.5)  
        try:

            # Create checkboxes for new items
            print ("try creating _checkbox")
            for i, item in enumerate(self.new_item_list_generate_doc):
                checkbox = QCheckBox(item)
                self.ui.scrollarea_checkbox_dirlist_000.addWidget(checkbox)
                self._checkboxes.append(checkbox)
        except Exception as e:
            print("Caught an exception: 3", e)
        
        try :
            # Delete checkboxes after the update process
            for checkbox in checkboxes_to_delete:
                checkbox.deleteLater()
        except Exception as e:
            print ( "Deleting checboxes for updating ", e)

        time.sleep (1)
        
        print ( "done creating checkbox ")
        QMessageBox.information(self, "Notice","New document added" )
        self._IS_ADDING_CHECKBOXES = False

        self.enable_show_g1 ()
        self.enable_show_ge2 ()
        self.enable_side_bar ()


    ## NOTE - selected items
    def _selected_items_for_generate_page(self):
        try:
            selected_items = [checkbox.text() for checkbox in self._checkboxes if checkbox.isChecked()]
            return selected_items
        except Exception as e:
            print ( "_selected_items_for_generate_page issue ", e )
            return []


    ## Function For Changing Page TO 
    ## User Page <<curently we do not need it>>
    def on_user_btn_clicked(self):
        print ( "do nothing on_user_btn_clicked" )


    ## Change QPushButton Checkable 
    ## Status When page_stackedWidget index changed
    def on_stackeWidget_currentChanged(self, index):
        btn_list    =   self.ui.icon_only_widget.findChildren(QPushButton)\
                        + self.ui.full_menu_widget.findChildren( QPushButton)
        
        for btn in btn_list:
            if index in [5,6]:
                btn.setAutoExclusive(False)
                btn.setChecked(False)
            else:
                btn.setAutoExclusive(True)


    ## SECTION - Settign Up 
    ## -------------------------------------------------------------------------->
    def setting_up_app(self):
        
        # for now we will use 2nd index of stack widget
        self.ui.page_stackedWidget.setCurrentIndex(1)

        # set QPushbutton( upload in first page ) text
        self.ui.pushButton_upload.setText("Chose File")

        # set up for search completer
        self.searh_completer = QCompleter ( self._search_path_list )
        self.searh_completer.setCaseSensitivity(0)
        self.ui.search_input.setCompleter ( self.searh_completer )

        # make_goup_one visible and enable
        self.update_progressbar (0)
        self.update_page_status ()
        self.hide_some_ui ()

    
    ## Group1 
    def enable_show_g1 ( self ):
        self.ui.pushButton_upload.setEnabled ( True )


    ## Group1 disable
    def disable_g1( self ):
        self.ui.pushButton_upload.setDisabled ( True )


    ## Group Edit 2 
    def enable_show_ge2( self ):
        self.ui.pushButton_update.setEnabled ( True )
        self.ui.pushButton_download.setEnabled ( True )
        self.ui.pushButton_remove.setEnabled ( True )
        self.ui.radioButton_newDoc.setEnabled ( True )
        self.ui.radioButton_addImage.setEnabled ( True)
        self.ui.search_btn.setEnabled ( True )
        self.ui.search_input.setEnabled (True)


    ## Group Edit 2 disable
    def disable_ge2( self ):
        self.ui.pushButton_update.setDisabled ( True )
        self.ui.pushButton_download.setDisabled ( True )
        self.ui.pushButton_remove.setDisabled ( True )
        self.ui.radioButton_newDoc.setDisabled ( True )
        self.ui.radioButton_addImage.setDisabled ( True )
        self.ui.search_btn.setDisabled ( True )
        self.ui.search_input.setDisabled ( True )

    
    ## Group side bar
    def enable_side_bar ( self ):
        self.ui.home_btn_1.setEnabled ( True )
        self.ui.home_btn_2.setEnabled ( True )
        self.ui.dashboard_btn_1.setEnabled ( True )
        self.ui.dashboard_btn_2.setEnabled ( True )
        self.ui.change_btn.setEnabled ( True )
        self.ui.pushButton_2_change_password.setEnabled ( True )
        self.ui.pushButton_3_change_password.setEnabled ( True )
        self.ui.pushButton_Change_lang.setEnabled ( True )
        self.ui.pushButton_2_change_language.setEnabled ( True )
        


    ## Group side bar disable
    def diasble_side_bar ( self ):
        self.ui.home_btn_1.setDisabled ( True )
        self.ui.home_btn_2.setDisabled ( True )
        self.ui.dashboard_btn_1.setDisabled ( True )
        self.ui.dashboard_btn_2.setDisabled ( True )
        self.ui.change_btn.setDisabled ( True )
        self.ui.pushButton_2_change_password.setDisabled ( True )
        self.ui.pushButton_3_change_password.setDisabled ( True )
        self.ui.pushButton_Change_lang.setDisabled ( True )
        self.ui.pushButton_2_change_language.setDisabled ( True )

    ## hide some elements
    def hide_some_ui ( self ):
        self.ui.radioButton_addImage.setHidden ( True )
        self.ui.radioButton_newDoc.setHidden ( True )
        self.ui.label_show.setHidden ( True )
        self.ui.input_area.setHidden ( True )



    ##----------------------------------------------------------------------------/>
    ## !SECTION


    ## SECTION - Pagination
    ## function for changing menu page ------------------------------------------->

    def on_home_btn_1_toggled(self):        ## for uplaodpage
        self._current_page_index = 0
        self.update_page_status ()
        # page menu setting
        if not self._IS_CREATING_MENU :
            self._IS_CREATING_MENU = True
            self.make_drop_down_menu()
    

    def on_home_btn_2_toggled(self):        ## for uplaodpage
        self._current_page_index = 0
        self.update_page_status ()
    

    def on_dashboard_btn_1_toggled(self):   ## for  download page
        self._current_page_index = 1
        self.update_page_status ()
        # page menu setting
        if not self._IS_CREATING_MENU :
            self._IS_CREATING_MENU = True
            self.make_drop_down_menu()
    

    def on_dashboard_btn_2_toggled(self):   ## downlaod page
        self._current_page_index = 1
        self.update_page_status ()
        self.ui.pushButton_3_generate
        self.ui.pushButton_2_genetate


    def on_pushButton_2_genetate_toggled ( self):
        try:
            self.generate_page ()
            
            # page menu setting
            if not self._IS_CREATING_MENU :
                self._IS_CREATING_MENU = True
                self.make_drop_down_menu()
        except Exception as e:
            print("Caught an exception:", e)
       
    
    def on_pushButton_3_genetate_toggled ( self):
        # self.generate_page ()
        print ("on_pushButton_3_genetate_toggled")
        

    def generate_page ( self ):
        # STUB - in this function we set variable to one but
        # page index to 0
        self._current_page_index = 1
        self.update_page_status ()
        self.ui.search_input.setDisabled ( True )
        self.ui.page_stackedWidget.setCurrentIndex (0)
        
        self._IS_NORMAL_PAGE = False
        
        # clear the checbox layout
        layout = self.ui.scrollarea_checkbox_dirlist_000.layout()
        if layout:
            while layout.count():
                item = layout.itemAt(0)
                widget = item.widget()
                if widget:
                    layout.removeWidget(widget)
                    widget.deleteLater()
                else:
                    layout.removeItem(item)


    def update_page_status( self ):
        if self._current_page_index == 0 :
            self.enable_show_g1 ()
            self.disable_ge2 ()
        else:
            self.enable_show_ge2 ()
            self.disable_g1 ()
            self.ui.pushButton_2_change_password
            self.ui.pushButton_3_change_password
            
        self.ui.page_stackedWidget.setCurrentIndex (1)
        self._IS_NORMAL_PAGE = True
        
       


    ##  --------------------------------------------------------------------------/>
    ## !SECTION - Pagination



    ## SECTION - Change password
    ## ---------------------------------------------------------------------------->
    def on_pushButton_2_change_password_pressed ( self):
        self.change_password()

    def on_pushButton_3_change_password_pressed ( self ):
        self.change_password()

    
    def change_password ( self ):
        self.ui.pushButton_2_change_password.setChecked ( True)
        self.ui.pushButton_3_change_password.setChecked ( True)

        password_ui = pass_w

        new_given_password = password_ui.set_password ()

        if not new_given_password ==  None and len ( new_given_password ) > 0:
            self.diasble_side_bar ()
            self.disable_g1 ()
            self.disable_ge2 ()
            self.update_progressbar (50)

            try:
                worker = WorkQT.Worker( self.change_password_procces, new_given_password )
                worker.signals.finished.connect ( self.change_password_finsihed )
                self.threadpool.start ( worker )
            except:
                traceback.print_exc ()

            

    def change_password_procces ( self,new_given_password , progress_callback ):

        try:
            password_check = CheckPassWord ()
            new_given_password = new_given_password.strip ()
            password_check.create_pass_at_pwd ( new_given_password )
            path = os.path.join ( os.path.dirname( os.path.abspath( __file__ ) ) , "pwd" )
            print( "change passwor dprocces path ", path)
            command = ["bypy", "delete", "/pwd/"]
            self.suproccess_show_plaintext ( command, progress_callback )
            time.sleep(2)
            command = [ "bypy", "upload", path , "/pwd/" ]
            self.suproccess_show_plaintext  ( command, progress_callback )
        except:
            traceback.print_exc ()

    
    def change_password_finsihed ( self ):
        self.ui.plainText_show.setPlainText ( "Password Changed 密碼已更改" )
        self.update_progressbar ( 100 )
        self.enable_side_bar ()
        self.enable_show_g1 ()
        self.enable_show_ge2 ()


    ## ----------------------------------------------------------------------------/>
    ## !SECTION - Change Passwor  


    ## SECTION - Change Language
    ## ----------------------------------------------------------------------------->

    def on_pushButton_Change_lang_pressed(self):
        self._is_english = not self._is_english
        self.update_language_setting ()
        


    def on_pushButton_2_change_language_pressed(self):
        self._is_english = not self._is_english
        self.update_language_setting ()


    def update_language_setting ( self ):
        print ( "update language ", self._is_english )
        if self._is_english:
            self.ui.home_btn_2.setText ( "Upload" )
            self.ui.dashboard_btn_2.setText ( "Operation")
            self.ui.pushButton_Change_lang.setText( "CN" )
            self.ui.exit_btn_2.setText ( "Exit" )
            self.ui.pushButton_update.setText ("Update")
            self.ui.pushButton_remove.setText ( "Remove" )
            self.ui.pushButton_download.setText ( "Download" )
            self.ui.pushButton_upload.setText ( "Upload" )
            self.ui.search_btn.setText ( "Search" )
            self.ui.pushButton_2_change_password.setText ( "Change Password" )
            self.ui.pushButton_3_generate.setText ( "Generate" )
            self.ui.pushButton_generate_doc.setText ( "Generate Doc")

        else :
            self.ui.home_btn_2.setText ( "上传" )
            self.ui.dashboard_btn_2.setText ( "服务")
            self.ui.pushButton_Change_lang.setText( "EN" )
            self.ui.exit_btn_2.setText ( "出口" )
            self.ui.pushButton_update.setText ("更新")
            self.ui.pushButton_remove.setText ( "消除" )
            self.ui.pushButton_download.setText ( "下载" )
            self.ui.pushButton_upload.setText ( "上传" )
            self.ui.search_btn.setText ( "搜索" )
            self.ui.pushButton_2_change_password.setText ("更改密碼")
            self.ui.pushButton_3_generate.setText ( "生成文档" )
            self.ui.pushButton_generate_doc.setText ( "生成文档" )



    ## -----------------------------------------------------------------------------/>
    ## !SECTION - Change Language




    ## SECTION - Remove Acton
    ## ---------------------------------------------------------------------------->

    ## when user clicked remove push button
    def on_pushButton_remove_pressed ( self ):

        # update progressbar
        self.update_progressbar ( 1 )

        # show message_box do they really want to delete
        dlg = QMessageBox(self)

        # setting language
        if self._is_english :
            dlg.setWindowTitle("Removing Contents")
            dlg.setText("Are  you sure you are willing to remove contents ??")
            dlg.setStandardButtons ( QMessageBox.Yes |  QMessageBox.No )
        else:
            dlg.setWindowTitle("删除内容")
            dlg.setText("您确定愿意删除内容吗 ??")
            dlg.setStandardButtons ( QMessageBox.Yes |  QMessageBox.No )
          
        

        button = dlg.exec ()
        # update progressbar
        self.update_progressbar ( 20 )


        if button == QMessageBox.Yes :
            dlg1 = QMessageBox(self)

            # setting language
            if self._is_english:
                dlg1.setWindowTitle("Removing Contents now")
                dlg1.setText("We will remove now ??")
                dlg1.setStandardButtons ( QMessageBox.Yes |  QMessageBox.No )
            else:
                dlg1.setWindowTitle("立即删除内容")
                dlg1.setText("我们现在将删除??")
                dlg1.setStandardButtons ( QMessageBox.Yes |  QMessageBox.No )
            
            button1 = dlg1.exec ()

            if button1 == QMessageBox.Yes :
                print ( "Removing Contents" )


                search_input = self.ui.search_input.text(  ).strip ( )
                if search_input :

                    print ( "user selected contents ", search_input )

                    # get value from search input
                    path = search_input.split( "  Directory:" ) [1]
                    print ( "Selected Path " , path)

                    worker = WorkQT.Worker (self.remove_proccess, path)
                    worker.signals.finished.connect ( self.removed )
                    self.threadpool.start ( worker )
                        
                
                else :
                    if self._is_english:
                        self.show_popup_text ( "File not selected", "First type text in search area then choose the file you want to make changes")
                    else :
                        self.show_popup_text ( "未选择文件", "首先在搜索区域中输入文本，然后选择要更改的文件")

            else:
                self.update_progressbar ( 100 )
        else:
            # update progressbar
            self.update_progressbar ( 100 )

    def remove_proccess(self, path , progress_callback):

        # check database is in use or not
        if self.database_is_in_use ():
            progress_callback.emit ("Database is in USE")
            return 0
        else:
            self.make_database_status_running ()


       

        # lets get path and the file name
        head, tail = os.path.split ( path )

        ondup = "ONDUP/" +path
        command = [ "bypy", "remove" , ondup ]
        self.suproccess_show_plaintext ( command, progress_callback)

        file_exists = self.subproccess_check_file_exists2 ( tail, head, progress_callback  )

      

        if file_exists :
            print ( " failled to remove " )
        else :
            try:
                self.open_database ()

                print ( Database.delete_path_list ( path ) )

                self.close_database ()

                pass
            except Exception as e:
                print ( " problem removing the list from database \'remove_proccess\' ",e )
      
        self.make_database_status_free ()


    def remove_progress (  self, s ):
        if "Database is in USE" in s:
            self.show_popup_text ( "Database is in USE 数据库正在使用中" )

    def removed (self):
        print ( "File removed")
        self.ui.plainText_show.setPlainText ( "内容删除成功(Removed content successful)" )
        
        self.clear_search_realted_variables()
        self.update_progressbar ( 100 ) ############## update progressbar
        
        self.disable_g1 ()


    

    ## later need to use worker signal to the background work

    ## ----------------------------------------------------------------------------/>
    ## !SECTION - Remove Action



    ## SECTION - Download Action
    ## ---------------------------------------------------------------------------->

    ## when user clicked remove push button
    def on_pushButton_download_pressed ( self ):

        search_input = self.ui.search_input.text(  ).strip ( )

        if search_input :

            self.update_progressbar ( 1 ) ############## update progressbar

            print ( "user selected contents ", search_input )

            # get value from search input
            path = search_input.split( "  Directory:" ) [1]
            print ( "Selected Path " , path)

            self.update_progressbar ( 3 ) ############## update progressbar

            # show folder opening option
            fname = QFileDialog.getExistingDirectory(self, 'Select Folder')


            self.update_progressbar ( 11 ) ############## update progressbar

            if fname:
                print ( "Selected path " , fname)
                worker = WorkQT.Worker (self.download_proccess, path, fname)
                worker.signals.progress.connect ( self.download_progress )
                worker.signals.finished.connect ( self.download_finished )
                self.threadpool.start ( worker )
                
            else:
                print ( "User did not chose anything " )
                self.update_progressbar ( 100 ) ############## update progressbar
        
        else :
            if self._is_english:
                self.show_popup_text ( "File not selected", "First type text in search area then choose the file you want to make changes")
            else :
                self.show_popup_text ( "未选择文件", "首先在搜索区域中输入文本，然后选择要更改的文件")

        self.ui.plainText_show.setPlainText ("(Process finished)处理完成")
        self.clear_search_realted_variables ()

        

    def download_proccess (self, bypy_path, local_path, progress_callback):
        bypy_path = "ONDUP/" + bypy_path
        command = [ "bypy", "download", bypy_path, local_path  ]
        self.suproccess_show_plaintext ( command, progress_callback)


    def download_progress (self, s):
        self.ui.plainText_show.setPlainText (s)

    
    def download_finished( self ):
        self.update_progressbar ( 100 ) ############## update progressbar
        self.disable_g1 ()

    ## later need to use worker signal to the background work

    ## ----------------------------------------------------------------------------/>
    ## !SECTION - Download Action



    ## SECTION - Update Action
    ## ----------------------------------------------------------------------------->
    ## when update button pressed
    def on_pushButton_update_pressed ( self ):

        # check the text vlaue in search input
        search_input = self.ui.search_input.text(  ).strip ( )
        
        if search_input and "Directory" in search_input :

            self.update_progressbar ( 1 ) ############## update progressbar

            print ( "user selected contents ", search_input )

            # get value from search input
            path = search_input.split( "  Directory:" ) [1]
            print ( "Selected Path " , path)


            try:

                self.open_database ()

                # get database value for giben *path*
                result = Database.get_info_based_on_path ( path )
                print ( " on_pushButton_update_pressed  ", str ( result ) )

                self.close_database ()

            except Exception as e :
                result = ["none"]
                print ("failled to insert for updating ", e)



            # new file name
            new_file_name = result[0][ "path" ].split ( "/" )[-1]
            print ( "File name ", new_file_name )

            # menu from json
            menu = result[0][ "sku_filter" ]
            print ( "Menu for update ", menu )


            self.update_progressbar ( 10 ) ############## update progressbar

            # ask user permission to open file
            dlg = QMessageBox(self)

            # setting language
            if self._is_english:
                dlg.setWindowTitle("Select File")
                dlg.setText("Select a  file to  upload")
            else:
                dlg.setWindowTitle("选择文件")
                dlg.setText("选择要上传的文件")

            self.update_progressbar ( 20 ) ############## update progressbar

            dlg.setStandardButtons ( QMessageBox.Yes |  QMessageBox.No )
            button = dlg.exec ()

            # if user accept show file opening dialog
            if button == QMessageBox.Yes :

                # ask user to select a file
                new_file_location_pc = QFileDialog.getOpenFileName(self, "Select file", "", "All Files (*)" )
                print ( "new_file_location_pc ", new_file_location_pc, " type ", type ( new_file_location_pc ) )
                new_file_location_pc = new_file_location_pc[0]
                print ( "new_file_location_pc ", new_file_location_pc, " type ", type ( new_file_location_pc ) )

                # get path for bypy from old file location  in bypy
                new_file_path_bypy = path.replace ( new_file_name, "" )
                print ( "new file path bypy", new_file_path_bypy )

                self.update_progressbar ( 30 ) ############## update progressbar


                ## NOTE let's match selected and old file name
                selected_file_extension = os.path.splitext( new_file_location_pc )[ 1 ]
                print ( "on_pushButton_update_pressed  selected_file_extension ", 
                       selected_file_extension )
                
                if selected_file_extension in new_file_name:
                    # connect worker 
                    worker = WorkQT.Worker ( self.update_ongoing_proccess ,
                                            new_file_location_pc , new_file_path_bypy, new_file_name, menu   )
                    worker.signals.progress.connect ( self.update_show_progress )
                    worker.signals.finished.connect ( self.update_complete )
                    

                    # Disable:: sidebar, Group 1-2
                    self.disable_g1 ()
                    self.disable_ge2 ()
                    self.diasble_side_bar () 

                    self.threadpool.start ( worker )

                else:
                    self.update_progressbar ( 100 ) ############## update progressbar
                    self.show_popup_text ( "Extension Error","Chosen file are not same type 選擇的文件類型不同" )
                    


            else:
                self.update_progressbar ( 100 ) ############## update progressbar
        else :
            if self._is_english:
                self.show_popup_text ( "File not selected", "First type text in search area then choose the file you want to make changes")
            else :
                self.show_popup_text ( "未选择文件", "首先在搜索区域中输入文本，然后选择要更改的文件")



    def update_ongoing_proccess (self, 
                                 new_file_location_pc , new_file_path_bypy, new_file_name, menu, progress_callback ) :
        
        

        if not len (  new_file_location_pc ) < 3:
             # check database is in use or not
            if self.database_is_in_use ():
                progress_callback.emit ("Database is in USE")
                return 0
            else:
                self.make_database_status_running ()
            
            print("Given file path ", new_file_location_pc)
            
            
            # copy file it to the temp folder
            print ( "Clear temp dir ", M_upload.clear_temp_dir())
            moduel_path  = M_upload.get_directory_path ()
            shutil.copy ( new_file_location_pc , moduel_path )


            # get selected file name
            head, tail = os.path.split ( new_file_location_pc )
            selected_file_name = tail

            # create a name for the file
            new_file_name_witout_extension = os.path.splitext ( new_file_name )[ 0 ]
            new_file_only_extension = os.path.splitext( new_file_name )[ 1 ]


            if "_copy(" in new_file_name_witout_extension :
                head, sep, tail2 = new_file_name_witout_extension.partition('_copy(')
                new_file_name_witout_extension = head


            now = datetime.datetime.now()
            now = now.strftime('%m-%d-%y_%H-%M')
            new_file_name_witout_extension = new_file_name_witout_extension + "_copy(" + str ( now )
            new_file_name = new_file_name_witout_extension + new_file_only_extension
            print ( "new file name  after adding extension and date ", new_file_name  )

            
            # rename document 
            M_upload.rename_file_in_temp ( selected_file_name, new_file_name )


            # get upload cmmand
            command = M_upload.get_bypy_upload_command  ( new_file_path_bypy )


            # NOTE - Update Document using subprocces
            self.suproccess_show_plaintext ( command, progress_callback )


            # check file upload succcesfull 
            file_uploaded = self.subproccess_check_file_exists ( new_file_name, progress_callback )


            if file_uploaded :
                new_file_path_bypy +=  new_file_name
                list1 =  new_file_path_bypy 
                try:
                    self.open_database ()

                    insert_db = Database.insert_dir_list( list1, menu )
                    print (" file uplaoded done , insert_db ", insert_db)

                    self.close_database ()

                except Exception as e :
                    print ( "failed to indert for update ", e )
                progress_callback.emit ( "(file upload done )文件上传完成" + str (new_file_path_bypy) )

            else:
                print ("failed to uplaod file")
                progress_callback.emit ( "(failled to upload file)文件上传失败" )


            # clear temp dir
            print ( M_upload.clear_temp_dir () )

        self.make_database_status_free ()


    # progress
    def update_show_progress(self, s):     
        self.ui.plainText_show.setPlainText ( s )   
        print (s)
        if "Database is in USE" in s:
            self.show_popup_text ( "Database is in USE 数据库正在使用中" )


    def update_complete(self ):
        print ( "Update done ")

        # Enable:: sidebar, group 2
        self.enable_show_ge2 ()
        self.enable_side_bar ()

        # clear search 
        self.clear_search_realted_variables ()

        self.update_progressbar ( 100 ) ############## update progressbar
        self.disable_g1 ()

        self.ui.plainText_show.setPlainText ("update done 更新完成")


    ## ----------------------------------------------------------------------------/>
    ## !SECTION - Update Action




    ## SECTION - UPLOAD button pressed
    ## when click upload button --------------------------------------------------->

    def on_pushButton_upload_pressed( self ):
        #
        self.Upload_Button_Clicked = True

        # disble upload button  and sidebar
        self.ui.pushButton_upload.setDisabled ( True )
        self.diasble_side_bar (  )

        # 
        self.show_dialog_chosing_file()
        

    ## NOTE - show dialog box
    # 
    def show_dialog_chosing_file( self ): 
        dlg = QMessageBox(self)
        
        # setting buttons for dialog
        # dlg.setStandardButtons(QMessageBox.Yes| QMessageBox.No)
        chose_folder = QPushButton()
        chose_file  = QPushButton()
        close_button = QPushButton()

        #  set text for different language
        if self._is_english:
            dlg.setWindowTitle("Opening File")
            dlg.setText("Here you have option to upload whole directory or" +
                        " you can just to upload a single file (.zip, .png, .jpg etc...) ")
            chose_file.setText(" Choose File" )
            chose_folder.setText( "Choose Folder" )
        else :
            dlg.setWindowTitle("打开文件")
            dlg.setText("在这里您可以选择上传整个目录，也可以只上传单个文件（.zip、.png、.jpg 等...）")
            chose_file.setText("选择文件" )
            chose_folder.setText( "选择文件夹" )


        chose_folder.clicked.connect(  self.folder_chosen )
        chose_file.clicked.connect(  self.file_chosen )
        close_button.clicked.connect ( self.file_chose_window_got_cancled )
        
        dlg.addButton ( chose_folder, QMessageBox.YesRole )
        dlg.addButton ( chose_file, QMessageBox.NoRole )
        # dlg.addButton ( close_button, QMessageBox.Cancel )

        # dlg.addButton ( )
        # dlg.setStandardButtons (  )
        dlg.setIcon(QMessageBox.Question)
        button = dlg.exec()

        if button != QMessageBox.YesRole or button != QMessageBox.NoAll :
            self.file_chose_window_got_cancled()


    ##  
    def file_chose_window_got_cancled ( self ):
        #
        self.Upload_Button_Clicked = False

        # disble upload button 
        self.ui.pushButton_upload.setEnabled ( True )
        self.enable_side_bar ()

    ##
    def folder_chosen(self):
        self._folder_chosen = True
        self.start_chosing()

    ##
    def file_chosen(self):
        self._folder_chosen = False
        self.start_chosing()

    ##
    ## Here we will start threadpool and other part
    def start_chosing(self):

        if self._folder_chosen:
            fname = QFileDialog.getExistingDirectory(self, 'Select Folder')
        else:
            fname = QFileDialog.getOpenFileName(self, "Select file", "", "All Files (*)" )
            

        # Open File Dialog
        # fname = QFileDialog.getOpenFileName(self, "Select Folder", "", "All Files (*)" )
		
		# Output filename to screen and pass it to upload module
        if fname:
            if self._folder_chosen:
                file_path   =   fname

                clear_dir = M_upload.clear_temp_dir() 
                print ("Clear fir")
                M_upload.copy_file_to_temp ( file_path )
            else:
                file_path = fname[0]
            print( " File selected ", file_path )

            worker = WorkQT.Worker(self.ongoing_proccess , file_path)
            worker.signals.progress.connect(self.show_progress)
            worker.signals.result.connect(self.show_result)
            worker.signals.finished.connect(self.thread_complete)

            time.sleep(0.5)
            
            self.update_progressbar ( 20 ) ############## update progressbar
            self.threadpool.start(worker)

            # test method for uplaod
            # self.test(file_path)


    ## Test methods for uplaod
    ##
    def unzip_upload_insert( self, path , menu , progress_callback ):

        print ("unzip_upload_insert \t\tMenu given ", menu, " \t path", path)
        if len(path) >5:
            # checking giving path belongs to a zip
            is_zip_file = M_upload.is_zip_file( path )

            print ( "unzip_upload_insert is_zip_file" , is_zip_file)
            if is_zip_file :
                # clearing temp folder
                clear_dir = M_upload.clear_temp_dir()
                print ( "clear_dir ", clear_dir["clear_temp_dir"] )
                progress_callback.emit ( "clear_dir"  )
                # unzip
                result = M_upload.start_unzipping ( path )

                print ( "unzip_upload_insert zipped", result )

                # if unzip succesfull
                if result[ "start_unzipping" ]:
                    print ( "Unzipped , will uplaod to baidu ")

                progress_callback.emit( "Unzipped" )

            else :
                if not self._folder_chosen:
                    M_upload.clear_temp_dir()
                    modue_path = M_upload.get_directory_path ()
                    shutil.copy( path, modue_path )
                    print ( "unzip inser folder chosen")

            progress_callback.emit( "uploading to baidu" )

            folder_list = M_upload.get_folder_list ()

            for item in folder_list :
                temp_file_name = ntpath.basename ( str(item) )
                # get selected file name
                new_file_name_witout_extension = os.path.splitext ( temp_file_name )[ 0 ]
                new_file_only_extension = os.path.splitext( temp_file_name )[ 1 ]

                if "_copy(" in new_file_name_witout_extension :
                    head, sep, tail2 = new_file_name_witout_extension.partition('_copy(')
                    new_file_name_witout_extension = head

                now = datetime.datetime.now()
                now = now.strftime('%m-%d-%y_%H-%M')
                new_file_name_witout_extension = new_file_name_witout_extension + "_copy(" + str ( now )
                new_file_name = new_file_name_witout_extension + new_file_only_extension
                print ( "new file name  after adding extension and date ", new_file_name  )

                progress_callback.emit ( "checking content exists or not 檢查內容是否存在 \t" +str ( new_file_name ) )

                if self.subproccess_check_file_exists (  temp_file_name, progress_callback ) or self.subproccess_check_file_exists (  new_file_name, progress_callback )  :
                    print (" unzip_upload_insert file name exists " , temp_file_name )
                    # rename document
                    M_upload.rename_file_in_temp ( temp_file_name, new_file_name )
                else:
                    print (" unzip_upload_insert file name does not exists ",temp_file_name  )



            folder_list = M_upload.get_folder_list ()
            print (" unzip_upload_insert ", folder_list)


            # # check same file name exists or not
            # file_exists= self.subproccess_check_file_exists ( new_file_name, progress_callback )

            # call upload module get the path
            command = M_upload.get_bypy_upload_command ()


            barv = 25
            barIndex = 0

            # use subproccess to uplaod
            self.suproccess_show_plaintext ( command, progress_callback)

            for item in folder_list :

                temp_file_name = ntpath.basename ( str(item) )
                progress_callback.emit( "uploading to baidu is done 上传至百度云盘完成 " +  str ( temp_file_name ) )

                head, sep, tail2 = new_file_name_witout_extension.partition('_copy(')
                new_file_name_witout_extension = head
                #  check file inserted or not
                if self.subproccess_check_file_exists ( temp_file_name, progress_callback ):
                    # insterting data
                    try :
                        self.open_database ()

                        insert_result = Database.insert_dir_list ( item, menu )
                        print ( "DB insert result" , insert_result )

                        self.close_database()
                    except Exception as e:
                        print ( "failed to insert for uplaod ", e )
            
        
        # clearing temp directory
        print ( " Clear temp folder ", M_upload.clear_temp_dir() )


        # enable upload button
        self.ui.pushButton_upload.setEnabled(True)

        # enable upload button and side bar
        self.Upload_Button_Clicked = False
        self.enable_side_bar (  )



    ## NOTE - CHeck database in use 
    #  
    def database_is_in_use ( self ):

        # first close the database
        self.close_database()

        # 2nd Download running.db and it's just downloading as
        #  database db also will be needed
        DB_MGM.download_db_file ()
        DB_MGM.download_db_running ()

        # 3rd reopen database connection
        self.open_database ()

        # 4th check running db whther user using them or not
        result, using_user_id = DatabaseRunning.is_database_in_use()

        # 5th close database connection
        self.close_database ()

        # now if database is in use check user id too,
        # if they are same user, then just say not in use 
        # Means, even if running db shwos status running , 
        # The return result will be false
        if result :
            using_user_id = using_user_id.strip ()
            if using_user_id == self._USER_ID:
                return False
        
        return result
    

    ## NOTE - Free Database status
    # also upload DB and DB_Running file
    def make_database_status_free ( self ):

        # 1st open database
        self.open_database ()

        # 2nd change the stutus free and user_id NONE
        DatabaseRunning.change_status ( "free", self._USER_ID )

        # 3rd close database
        self.close_database ()

        # 4th upload it to the baidu
        DB_MGM.upload_db_running ()
        DB_MGM.upload_db_file ()

    
    ## NOTE - Free Database status
    # also upload DB and DB_Running file
    def make_database_status_running ( self ):

        # 1st open database
        self.open_database ()

        # 2nd change the stutus free and user_id NONE
        DatabaseRunning.change_status ( "running", self._USER_ID )

        # 3rd close database
        self.close_database ()

        # 4th upload it to the baidu
        DB_MGM.upload_db_running ()
        DB_MGM.upload_db_file ()

    
    ## SECTION  WORKER Class method For Upload
    def ongoing_proccess(self,file_path,progress_callback):  
        # show thar process starte
        progress_callback.emit("Process started")


        # first check all the comboxes selected or not
        all_comboboxes_selected = False
        all_text = ""

        # check all existing comboboxes in scroll_area_editpage
        i = 0
        for menu in self.combo_box_list :
            text = menu.currentText ()

            if text :
                print ( "Got selected menu ", text )
                all_comboboxes_selected = True
                if i > 0:
                    all_text += "." + text
                else:
                    all_text +=  text
            else :
                print ( "Combobox not selected" )
                all_comboboxes_selected = False
            i += 1

        # if comboboxes are not all set show message
        if not all_comboboxes_selected :
            progress_callback.emit("menu not selected")


        else :

            # chec whether database is in use
            if self.database_is_in_use ():
                progress_callback.emit ( "Database is in USE" )
                return None
            
            # now change the database status
            else :
                self.make_database_status_running ()
            try:
                self.unzip_upload_insert ( file_path, all_text, progress_callback )
            except Exception as e :
                print ( "self.unzip_upload_insert ", e )

            # free up using the database
            self.make_database_status_free ()

        return "Finished Uploading"

    ## show progress
    def show_progress(self, s):
        self.ui.plainText_show.setPlainText ( s )

        # if menu not selected
        if "menu not selected" in s :
            if self._is_english:
                self.show_popup_text ( "Menu missing !!", " All menu not selected, please select all menu first then click upload" )
            else :
                self.show_popup_text ( "菜单缺失！！","所有菜单未选择，请先选择所有菜单然后点击上传" )
        elif "Database is in USE" in s:
            self.show_popup_text ( "Database is in USE 数据库正在使用中" )
    ## result
    def show_result(self, s):
        print("Showing resutl done for this even ",s)

    ## thread complete
    def thread_complete(self):
        print("Thread for uplaod done")

        self.ui.plainText_show.setPlainText ( "Upload Completed 上传完成" )

        self.update_progressbar ( 100 ) ############## update progressbar

        self.disable_ge2 ()



    # END--------------------------------------------------------------------------/>
    ## !SECTION worker methods for Upload 



    ## END-------------------------------------------------------------------------/>
    ## !SECTION Upload Button



    ##  SECTION - Make Drop Down Menu
    # -------------------------------------------------------------------------------> 
    def make_drop_down_menu(self, page_index=0):

        self._selected_menu_json_path = None
        self._initial_menu_list = None
        self._Menu_Json = None
        self._first_key = ""

        self.vbox_menu = QVBoxLayout()
        self.combo_box_list : QComboBox() = [] 

        # test ::=>
        #   json_value = Read_Write_Menu.get_json_for_menu()
        #   print(" Json value from the md fiel ", json_value)
        #   Read_Write_Menu.create_menu_json_file(json_value)
        
        self.remove_All_the_elements ()
        worker = WorkQT.Worker(self.on_going_process_for_initial_menu_creation)
        worker.signals.result.connect(self.result_for_initial_menu_creation)
        self.threadpool.start(worker)

        #NOTE - making menu path empty
        self._selected_menu_json_path = ""
        self.combo_box_list = []


    def remove_All_the_elements ( self ):
        # Assuming self.scroll_content_0 is your QScrollArea's widget
        content_widget_0 = self.scroll_content_0.widget()

        # Clear the layout of the content widget
        content_layout_0 = content_widget_0.layout()
        if content_layout_0:
            while content_layout_0.count():
                item = content_layout_0.itemAt(0)
                widget = item.widget()
                if widget:
                    content_layout_0.removeWidget(widget)
                    widget.deleteLater()
                else:
                    content_layout_0.removeItem(item)

        # Repeat the same steps for self.scroll_content_generate
        content_widget_generate = self.scroll_content_generate.widget()
        content_layout_generate = content_widget_generate.layout()
        if content_layout_generate:
            while content_layout_generate.count():
                item = content_layout_generate.itemAt(0)
                widget = item.widget()
                if widget:
                    content_layout_generate.removeWidget(widget)
                    widget.deleteLater()
                else:
                    content_layout_generate.removeItem(item)

    def on_going_process_for_initial_menu_creation(self, progress_callback):
        progress_callback.emit ( " Reading Menu From 'menu.md' file " )
        self._Menu_Json = Read_Write_Menu.get_json_for_menu ()
        progress_callback.emit ( "Reading done successfuly " )

        self._first_key = list ( self._Menu_Json.keys() )[0]

        self._initial_menu_list = list (self.get_list_from_given_data(self._Menu_Json))

        # NOTE - setting initial menu path
        self._selected_menu_json_path += self._first_key 


    def result_for_initial_menu_creation(self, s):
        label = QLabel(self)
        label.setText("Choose File 选择文件" )
        

        combo = QComboBox(self)
        combo.addItems(self._initial_menu_list)
        combo.setCurrentIndex(-1)

        layout = QHBoxLayout()
        layout.addWidget(label)
        layout.addWidget(combo)

        widget = QWidget()
        widget.setFixedHeight(40)
        widget.setLayout(layout)

        self.vbox_menu.addWidget(widget)

        widget = QWidget()
        widget.setLayout(self.vbox_menu)

        if self._IS_NORMAL_PAGE:
            self.scroll_content_0.setWidget( widget )
        else:    
            self.scroll_content_generate.setWidget ( widget )

        self.combo_box_list.append(combo)

        temp_lenght= len(self.combo_box_list)

        self.combo_box_list[ temp_lenght -1 ].activated[int].connect( 
            partial(self.try_to_add_next_menu,   temp_lenght )) 
        
        self._IS_CREATING_MENU = False
        
        
    def try_to_add_next_menu(self, next_menu_index):

        current_menu_text = ""

        try :

            # If we are getting a value here, means user pressed a previous menu button
            current_menu_text = self.combo_box_list[ next_menu_index ]

            # index will be -1 
            index = next_menu_index -1

            # Based on current menu index we will do the deletion of menu
            print(" Menu combo box selected is  ", index)

            # We need to reset few thing:
            #                        1.  selected_json_path
            #                        2.  combox_list
            #                        3.  VBOX_layout
            #                        4.  Add value to last indexed list

            # reset combo_list 
            temp_combo_list = self.reset_combo_list(index)

            # add values to the last indexed combo list 
            self.reset_comboxes_item()

            # reset selected_json_path
            temp_p = self.reset_selected_json_path(index) 

            self.show_curren_selected_json_path("Try, where user previous click got clicked")

            # reset VBOX layout 
            self.reset_v_box_layout(self.vbox_menu, index)

           
            
        except Exception as e:
            print("Next Menu Does not exits or other problem" , e)
            print(traceback.format_exc())

            # lets check our menu list size
            current_menu_size = len ( self.combo_box_list)

            # next menu list does not exist
            # so lenght of  menu list will be same as next_menu_index
            if( next_menu_index == current_menu_size):


                #
                #***Another scenerio where lenght and menu_index can be same 
                # The JSON list is already at the end
                # currently it's does not create any problem
                # So we can let it be as it is

                # NOTE - updating menu path 
                self.update_menu_path()
                 
                # self.show_curren_selected_json_path("\\try_to_add_menu\\")
                
                #
                # creating menu
                # first create a list for current path
                temp_path_list = self._selected_menu_json_path.split(".")


                try:
                    # The initial key is not present , we take it in advance
                    temp_json = self._Menu_Json[temp_path_list[0]]

                    # then based on current trail we will make new menu
                    for last_key in temp_path_list[1:]:
                        temp_json = temp_json [last_key]
                    self.add_children_menu_to_the_ui(temp_json)

                except:
                    path = self._first_key+ "." + str( self.combo_box_list[0].currentText())
                
                    for x in self.combo_box_list[1:]:
                        path +="." + str( x.currentText())
                    self._selected_menu_json_path = path
                    print("path for ", self._selected_menu_json_path)
            

            #  
            # Now maybe user clicked Menu which is not last branch
            # let we need to check if selected JSON menu has more list
            # Even we do not check it will work
            else:
                print("we are here exception did not handle some issues"  )


    def add_children_menu_to_the_ui(self, item):
       
        combo = QComboBox(self)
        combo.addItems(item)
        combo.setCurrentIndex(-1)

        layout = QHBoxLayout()
        layout.addWidget(combo)

        widget = QWidget()
        widget.setFixedHeight(40)
        widget.setLayout(layout)

        self.vbox_menu.addWidget(widget)

        self.combo_box_list.append(combo)

        temp_lenght= len(self.combo_box_list)
      
        self.combo_box_list[ temp_lenght -1 ].activated[int].connect(partial(self.try_to_add_next_menu,   temp_lenght ))

        self.show_curren_selected_json_path("\\add_children_menu_to_the_ui\\")


    def get_list_from_given_data(self, value ) :
        key = list(value.keys())[0]
        return value[key].keys()


    # adding new value when combo box  is slected 
    # only call when user selected the current menu
    def update_menu_path(self):
        self._selected_menu_json_path = self._first_key
        for item in self.combo_box_list:
            self._selected_menu_json_path += "." + item.currentText()


    # reseting combo list
    # we will get index of curent selected menu 
    # need to add 1, 
    # which  means from 0th positon, 
    # we will take n(index)th number of element
    # we need to check, the "KEY" has list or not
    def reset_combo_list( self, index ):

        add_item = 1

        # before adding two lets confirm, the current menu has submen or not
        self.combo_box_list = self.combo_box_list[ 0 : index + 2 ]

        self.combo_box_list[-1].clear()
        
        return self._selected_menu_json_path


    # reset selected_json_path
    # here one extra key always added from the begining
    # 
    def reset_selected_json_path( self, current_menu_index):
        temp = self._first_key
        for item in self.combo_box_list:
            temp += "." + item.currentText()
        self._selected_menu_json_path = temp

        return temp


    # resetting Vbox layout
    def reset_v_box_layout( self,layout, index ):
        print ("Vbox count ", layout.count() , " len of combo boxes ", len(self.combo_box_list))
        
        for i in reversed( range( layout.count() ) ): 
                if layout.count() <= len(self.combo_box_list):
                    break
                layout.itemAt( i ).widget().deleteLater()
                layout.itemAt( i ).widget().setParent(None)
                

    # reseting comboboxes items at -1th position
    def reset_comboxes_item( self  ):

        i = 0
        temp_json = self._Menu_Json[ self._first_key ] 

        for item in self.combo_box_list:
            if item.currentText():
                print("Comboxes text ", item.currentText() )
                temp_json = temp_json[ item.currentText() ]
                pass
            else :
                if isinstance(temp_json, dict ): 
                    print( "Curent index does not have text, so will add dic \n\t\t ",
                        " ", temp_json.keys() )
                    
                    new_list = list ( temp_json.keys() )
                    self.combo_box_list[-1].addItems(new_list)
                    self.combo_box_list[-1].setCurrentIndex(-1)
                elif isinstance( temp_json, str ):
                    print( "Curent index does not have text, so will add str \n\t\t ",
                        " ", temp_json, "\t size", len(temp_json)  )
                    
                    self.combo_box_list[-1].addItem(temp_json)
                    self.combo_box_list[-1].setCurrentIndex(-1)

                    if (len (temp_json) == 0 ):
                        print (" Size of list ", len (self.combo_box_list) )
                        self.combo_box_list = self.combo_box_list[:-1]
                        print("poping list ")
                        print (" Size of list ", len (self.combo_box_list) )

                elif isinstance( temp_json, list ):
                    print( "Curent index does not have text, so will add list \n\t\t ",
                        " ", temp_json )
                    self.combo_box_list[-1].addItems(temp_json)
                    self.combo_box_list[-1].setCurrentIndex(-1)

    
    # chec selected menu has sub menu or not
    def selected_has_sub_menu(self, text, index):
        temp_json = self._Menu_Json[ self._first_key ] 

    def show_curren_selected_json_path(self, from_where=None):
        print("Path " , self._selected_menu_json_path , "   From ", from_where 
              , "\t lenght of current menu list ", len(self.combo_box_list))
        
    ## END----------------------------------------------------------------------------/>
    ## !SECTION MakeDrop Down Menu


    ## SECTION Search Result 
    #  -------------------------------------------------------------------------------->
    # Update search path list
    # This method will take value for "self._search_path_list"
    # It will just add them to the completer


    def update_search_path( self ):
        model = self.searh_completer.model ()
        model.setStringList( self._search_path_list)
        # print ("Search completer updated ")
        

    # get search list from 
    def get_search_path_list_from_db ( self, progress_callback ):

        # Comobox nth menu for now by deafult None
        # if loop through all menu return no text , it will be none
        # we will get from another method
        filter_menu = ""
        for item in self.combo_box_list:
            if item.currentText () is not None:
                filter_menu  += "." + item.currentText ()

        # Try to get text from search
        # searc_txt = self.ui.search_input

        # call database to get data on this 
        try :
            self.open_database ()

            result_list = Database.get_dir_list( filter_menu )
            # print ("result on change text search input ", result_list)

            self.close_database ()
        except Exception as e :
            result_list = []
            print ( "get_search_path_list_from_db ", e )
        
        # print ( result_list )

        # put the list in self._search_path_list
        self._search_path_list = result_list

        

    def search_finsihed ( self, s):
        self.update_search_path()

    def on_search_input_textChanged ( self ):
        print ("search input text ", self.ui.search_input.text())

        # worker = WorkQT.Worker ( self.get_search_path_list_from_db )
        # worker.signals.finished.connect ( self.search_finsihed )
        # self.threadpool.start ( worker )

        # can delete later
        filter_menu = ""
        for item in self.combo_box_list:
            if item.currentText () is not None:
                filter_menu  += "." + item.currentText ()

        # Try to get text from search
        # searc_txt = self.ui.search_input

        # call database to get data on this 
        try :
            self.open_database ()

            result_list = Database.get_dir_list( filter_menu )

            self.close_database()

        except Exception as e :
            result_list = []
            print ( "on_search_input_textChanged ", e )
        # print ( result_list )

        # put the list in self._search_path_list
        # print ( "search result ", result_list )
        self._search_path_list = result_list
        self.update_search_path()

    #  -------------------------------------------------------------------------------/>
    ## !SECTION





    ## SECTION - for generating document
    #  -------------------------------------------------------------------------------->
    #
    ## SECTION -  arrow button
    # on clicked on the arrow button
    def on_button_add_list_to_generate_pressed ( self ):
        print ( "Arrow Button Pressed", self._selected_items_for_generate_page () )
        
        # update list
        self.update_list_for_document_generation_left_Side ()



    # NOTE -  left side list adding
    def update_list_for_document_generation_left_Side ( self ):
        # try list to ListWidget
        try:
            # save value inside generate doc list
            self.add_new_values_to_generate_doc_list ( self._selected_items_for_generate_page () )

            # make the ListWidget clear first
            self.ui.listWidget.clear ()

            # adding "self.generate_doc_list" list to the ListWidget
            self.ui.listWidget.addItems ( self.generate_doc_list   )

            # adding drag and reorder options
            self.ui.listWidget.setDragDropMode(QListWidget.InternalMove)

        except Exception as e:
            print ( "Problem creating list for generating page ", e )



    # functions to stores only new path to the "self.generate_doc_list"
    def add_new_values_to_generate_doc_list(self, new_values_list):
        print ( "given list in \"add_new_values_to_generate_doc_list\"", new_values_list )
        try:
            for value in new_values_list:
                if value not in self.generate_doc_list:
                    self.generate_doc_list.append(value)
        except Exception as e:
            print (" trying to add values to the left side  ", e)

    ## !SECTION -  arrow button



    ## SECTION -  CROSS BUTTON 
    def on_button_clear_list_to_generate_pressed ( self ):
        print ( "Cross button pressed  selected items " )
        try:
            # make the ListWidget clear first
            self.ui.listWidget.clear ()
            self.generate_doc_list.clear ()
        except Exception as e:
            print (" trying to remove list  from left side ", e)
            

    ## !SECTION -  CROSS BUTTON 



    ## SECTION -  Generate Document
    #  

    def on_pushButton_generate_doc_pressed ( self ):
        print ( "Generate doc button pressed " )

        # get item from left side
        doc_items = self.get_list_from_left_Side ()

        # only get the directory
        doc_items = self.extract_filenames_with_directory_left_Side ( doc_items )

        # get new doc name
        user_given_doc_name = self.get_new_genrated_doc_name ()

        # Prompt the user to choose an output folder
        output_folder = QFileDialog.getExistingDirectory(self, "Select Output Folder")
        name = self.ui.lineEdit_new_doc_name.text ().strip ()
        if output_folder and len (name) >1:
                print("Got folder location:", output_folder)
                # download all from baidu
                self.download_files_from_baidu_to_temp ( doc_items, output_folder )
        else :
            warning_text = "No output folder selected or no file name provided\n未选择输出文件夹。 或未提供文件名"
            QMessageBox.warning(self, "Warning / 警告", warning_text)

        


    ## SECTION - worker class to download
    # show current progress, starting function and proccessed finished one
    # 
    def download_files_from_baidu_to_temp ( self,  baidu_file_list, output_folder ):

        self.progress_dialog = ProgressDialog(self, "Generating Document pls wait")
        self.progress_dialog.show()

        worker = WorkQT.Worker ( self.generate_document_ongoing_proccess, baidu_file_list, output_folder )
        worker.signals.progress.connect ( self.generate_document_ongoing_progress )
        worker.signals.finished.connect ( self.generate_document_finsihed )

        self.threadpool.start ( worker )


    # let's check the downloaded files
    def check_download_status_for_generating_page ( self, downloaded_files, path):
        # Get the list of base names of downloaded files
        downloaded_base_names = [os.path.basename(downloaded_file) for downloaded_file in downloaded_files]
        
        # print ( "check_download_status_for_generating_page ",  downloaded_base_names, "\nfor given file list ", downloaded_files )
        # Get the list of base names of existing files in the path
        existing_files = [os.path.basename(existing_file) for existing_file in os.listdir(path)]
        
        # Check if all downloaded files exist in the existing files
        all_files_downloaded = all(downloaded_base_name in existing_files for downloaded_base_name in downloaded_base_names)
        
        # Create a list of missing files
        missing_files = [downloaded_base_name for downloaded_base_name in downloaded_base_names if downloaded_base_name not in existing_files]
        
        # Set the status based on whether all files are downloaded
        status = "ok" if all_files_downloaded else "error"
        
        # Create the response dictionary
        response = {
            "all_file_downloaded": all_files_downloaded,
            "missing_files": missing_files,
            "status": status
        }
        
        return response
        

    # let's show an window
    def show_download_status_for_generating_page ( self,  response ):
        message_box = QMessageBox( self )
        message_box.setWindowTitle("Download Status")
        
        if response["all_file_downloaded"]:
            message_box.setIcon(QMessageBox.Information)
            message_box.setText("All files downloaded successfully.")
        else:
            message_box.setIcon(QMessageBox.Warning)
            missing_files = "\n ".join(response["missing_files"])
            message_box.setText(f"Some files were not downloaded: {missing_files}")
        
        message_box.setStandardButtons(QMessageBox.Ok)
        message_box.show()


    # now lets create a file for selected documents
    def add_files_to_doc( self, file_list, folder_path, output_filename, output_folder):
        doc = Document()
        composer = Composer(doc)

        print("File list ", file_list)

        # Set the standard height for images
        standard_height = 300  # You can adjust this value

        for filename in file_list:
            file_path = os.path.join(folder_path, filename)
            print(" add_files_to_doc file path ", file_path)

            try:
                if filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                    img = Image.open(file_path)
                    img = img.convert('RGB')

                    # Set the desired aspect ratio (height-to-width ratio)
                    desired_aspect_ratio = 0.75  # You can adjust this value

                    # Calculate the new width based on the standard height and desired aspect ratio
                    new_width = int(standard_height / desired_aspect_ratio)

                    # Resize the image to the new dimensions (standard_height x new_width)
                    img = img.resize((new_width, standard_height), Image.ANTIALIAS)

                    temp_image_path = os.path.join(output_folder, 'temp_image.jpg')
                    img.save(temp_image_path, format='JPEG')
                    doc.add_picture(temp_image_path)
                    os.remove(temp_image_path)

                elif filename.lower().endswith(('.doc', '.docx')):
                    docx_content = Document(file_path)
                    composer.append(docx_content)

            except Exception as e:
                print(f"Error processing {filename}: {e}")

        output_path = os.path.join(output_folder, output_filename)
        composer.save(output_path)
        print(f"Document '{output_filename}' saved successfully in '{output_folder}'.")



    # 
    def generate_document_ongoing_proccess ( self,  baidu_file_list, output_folder, progress_callback ):

        # Check for unsupported extensions
        unsupported_extensions = ['jpg', 'jpeg', 'png', 'doc', 'docx']
        unsupported_files = [file for file in baidu_file_list if not any(file.endswith(ext) for ext in unsupported_extensions)]

        if unsupported_files:
            progress_callback.emit("extensionerror")
            return
        #download folder path 
        path = M_upload.get_directory_path ( )
        print ( "path to download ", path  )

        # clear dir
        result = M_upload.clear_temp_dir ( )
        print ( "Directory is ready for left side list downlaod ", result )

        # download files
        try:
            self.setEnabled(False)  # Disable the main window
            total_files = len(baidu_file_list)
            for index, baidu_file in enumerate(baidu_file_list):
                # Calculate the progress percentage
                progress_percentage = int((index + 1) / total_files * 100)
                # Combine the progress percentage
                progress_info = f"progress::{progress_percentage}"
                # Emit progress callback with the progress information
                progress_callback.emit(progress_info)

                baidu_file = "/ONDUP/"+ baidu_file
                
                # Download the file using bypy
                bp.download(baidu_file, path)
                print(f"Downloaded file {index + 1}/{total_files}: {baidu_file}")
            
            print("Files downloaded successfully.")
            time.sleep (1)
        except Exception as e:
            
            self.setEnabled(True)  # Enable the main window
            print("An error occurred:", e)


        try:
            result = self.check_download_status_for_generating_page ( baidu_file_list, path )
            print ( "rsult ", result )
            if result.get("all_file_downloaded"):
                print ( "all file downlaoded" )
            else:
                print ( "all file are not downlaoded" )
                response_string = str(result)
                progress_callback.emit(response_string)
        except Exception as e:
            
            self.setEnabled(True)  # Enable the main window
            print("An error occurred:", e)

        if result.get("all_file_downloaded"):
            time.sleep (3)
            file_list = [os.path.basename(path) for path in baidu_file_list]
            
            folder_path = path
            name = self.ui.lineEdit_new_doc_name.text ()
            name = name.strip ()

            if len (name) > 1 :
                output_filename = name + ".doc"
            else:
                output_filename = 'output.doc'
            
            
            self.add_files_to_doc ( file_list, folder_path, output_filename, output_folder )
            
            
            M_upload.clear_temp_dir ()


    
    #
    def generate_document_ongoing_progress(self, progress_response):
        if progress_response.startswith("progress::"):
            # Extract the percentage from the progress string
            progress_percentage = int(progress_response.split("::")[-1])
            # Update the progress bar with the extracted percentage
            self.progress_dialog.progress_bar.setValue(progress_percentage)
            print("from \"generate_document_ongoing_progress\" Progress:", progress_percentage)
        elif progress_response == "extensionerror":
            # Handle extension error
            message_box = QMessageBox( self )
            message_box.setWindowTitle("Extension Error")
            message_box.setIcon(QMessageBox.Warning)
            message_box.setText("Only files with extensions: jpg, jpeg, png, doc, docx are supported.")
            message_box.setStandardButtons(QMessageBox.Ok)
            message_box.show()

            self.progress_dialog.close()
            self.setEnabled(True)  # Enable the main window
        else:
            response = eval(progress_response)
            if response.get("all_file_downloaded"):
                print("All files downloaded successfully.")
            else:
                print("Some files were not downloaded:", ", ".join(response["missing_files"]))

            self.progress_dialog.close()
            self.setEnabled(True)  # Enable the main window
            self.show_download_status_for_generating_page(response)


    #
    def generate_document_finsihed( self ):
        self.setEnabled(True)  # Enable the main window
        self.progress_dialog.close()
        print ( "done generating new document" )

    #
    ##  
    ## !SECTION - 

    def get_list_from_left_Side ( self ):
        try:
            ordered_items = [self.ui.listWidget.item(i).text() for i in range(self.ui.listWidget.count())]
            print("Pressed Generated Button :: reordered list", ordered_items)
            return ordered_items
        except Exception as e:
            print("An error occurred:", e)
            return []


    def get_new_genrated_doc_name ( self ):
        try:
            text = self.ui.lineEdit_new_doc_name.text ()
            print (  "Pressed Genrated Button :: new doc name ", text  )
            return text
        except Exception as e:
            print("An error occurred:", e)
            return None

    

    def extract_filenames_with_directory_left_Side( self,  data_list ):
        try:
            extracted_filenames = [item.split('Directory:')[-1].strip() for item in data_list if 'Directory:' in item]
            return extracted_filenames
        except Exception as e:
            print("Error:", e)
            return []

    ## !SECTION -  Generate Document
    #
    #
    #  -------------------------------------------------------------------------------/>
    ## !SECTION - for generaitng documnet




    ## SECTION - common methods for UI
    ## --------------------------------------------------------------------------------->

    def change_button_text( self, widget , text ):
        widget.setText(text)


    def change_button_state( self, widget, text ):
        pass


    # here state false means, disable
    def enable_button( self,  widget, state ):

        if not state:
            widget.setDisabled()
        else:
            widget.setEnabled()

    # clear search area 
    def clear_search_realted_variables (self):
        self._search_path_list = []
        self.update_search_path ()
        self.ui.search_input.clear ()


    # progress bar showing , hide and bar changing
    def update_progressbar( self, percentage_done ):

        self.ui.progressBar.setValue ( percentage_done )

        if percentage_done == 0 or percentage_done == 100:
            self.enable_show_g1 ()
            self.enable_show_ge2 ()
            self.enable_side_bar()
            time.sleep (2)
            self.ui.progressBar.hide ()


        elif percentage_done >0 and percentage_done < 100:
            self.ui.progressBar.setVisible ( True )
            time.sleep (2)
            self.disable_g1 ()
            self.disable_ge2()
            self.diasble_side_bar()



    # show message box without tex
    def show_popup_text ( self , title, details ) :
        # show message_box do they really want to delete
        dlg = QMessageBox( self )
        dlg.setWindowTitle( title )
        dlg.setText( details )
        dlg.setStandardButtons ( QMessageBox.Ok )
        dlg.exec ()

    ## NOTE - Close database
    def close_database ( self ):
        try:
            Database.close_db_connection()
            DatabaseRunning.close_connection ()
        except:
            traceback.print_exc ()

    ## NOTE - Open database
    def open_database ( self ):
        try:
            Database.open_db_connection()
            DatabaseRunning.open_connection()
        except:
            traceback.print_exc ()

  

    # NOTE - using subproccessing to do bypy work 
    def suproccess_show_plaintext ( self, command_list , progress_callback )  :

        try:
            proc = subprocess.Popen(command_list, stdout=subprocess.PIPE)

            while True:
                line = proc.stdout.readline()
                if not line:
                    break

                               # print ( "test:", line.rstrip() )
                try:
                    value = str ( line, "utf-8") .strip()
                    print("subprocces ", value)
                    progress_callback.emit ( value )
                except:
                    progress_callback.emit ( line )
                    traceback.print_exc ()
                
        except :
            traceback.print_exc()



    
    # NOTE - use subprocces  to check file exists or not
    def subproccess_check_file_exists ( self, file_name, progress_callback  ):
        try:
            text_Found = False
            file_name_exists = False

            print (" File name from subprocces ", file_name )
            command = [ "bypy", "search", file_name ]
            proc =  subprocess.Popen( command  , stdout=subprocess.PIPE )
            while True:
                line = proc.stdout.readline()
                if not line:
                    break

                # print ( "test:", line.rstrip() )
                try:
                    value = str ( line, "utf-8") .strip()
                    print("subprocces ", value)
                    

                    progress_callback.emit ( value )
                    if "Found" in value :
                        text_Found = True
                    if text_Found and file_name in value:
                        file_name_exists = True
                except:
                    value = bytes (line)
                    line = str ( line )
                    print ( line )
                    #value = ( line ).encode('utf-8').decode('unicode_escape')
                    print ( "*"*10)
                    print ( value )
                    print ( "*"*10)
                    if text_Found :
                        file_name_exists = True
                    
                    progress_callback.emit ( line )
                    traceback.print_exc ()

                

            if file_name_exists:
                return True
        except:
            print(traceback.format_exc())
            return  None
        
            

    def subproccess_check_file_exists2 ( self, file_name, remote_path ,progress_callback  ):
        try:
            text_Found = False
            file_name_exists = False

            print (" File name from subprocces ", file_name, "\t remote path ", remote_path )
            command = [ "bypy", "search", file_name, remote_path ]
            proc =  subprocess.Popen( command  , stdout=subprocess.PIPE )
            while True:
                line = proc.stdout.readline()
                if not line:
                    break

                # print ( "test:", line.rstrip() )
                try:
                    value = str ( line, "utf-8") .strip()
                    print("subprocces ", value)
                    

                    progress_callback.emit ( value )
                    if "Found" in value :
                        text_Found = True
                    if text_Found and file_name in value:
                        file_name_exists = True
                except:
                    value = bytes (line)
                    line = str ( line )
                    print ( line )
                    #value = ( line ).encode('utf-8').decode('unicode_escape')
                    print ( "*"*10)
                    print ( value )
                    print ( "*"*10)
                    if text_Found :
                        file_name_exists = True
                    
                    progress_callback.emit ( line )
                    traceback.print_exc ()
        except:
            print(traceback.format_exc())
            return  None

    ## ---------------------------------------------------------------------------------/>
    ## !SECTION

"""
 #NOTE END-Main Window---------------------------------Main Window------------------------------------------>      
"""

#SECTION - password
class Password ( QMainWindow ):

    def __init__(self):
        super().__init__( )


    def get_password (self):
        dlgi = QInputDialog()
        text, ok = dlgi.getText(self, '密码(password)', ' ')
        dlgi.closeEvent = self.CloseEvent
        self.show()

        if ok:
            print( "User given text", text )
            self.hide()
            return text
        else :
            self.hide()
            return None

    def CloseEvent(self, event):
        print("clossed ")
        return None
        
    def set_password (self):
        dlgi = QInputDialog()
        text, ok = dlgi.getText(self, '輸入新密碼(input password)', ' ')
        self.show()
        if ok:
            # print( "User Given Text", text )
            self.hide()
            return text
        else :
            self.hide()
            return None
        

class CheckPassWord() :
    def __init__(self) :
        self.base_path = M_upload.get_directory_path ()
        self.pwsd_dir = "pwd"


    def create_pwd_folder ( self ):
        try:
            print ( "base directory ", self.base_path )

            self.pwsd_dir = os.path.join ( os.path.dirname( os.path.abspath( __file__ ) ) , self.pwsd_dir )
            print ( "psw directory ", self.pwsd_dir )
            os.mkdir (self.pwsd_dir)    
        except:
            traceback.print_exc()
    

    def remove_pwd_folder ( self ):
        try:
            self.pwsd_dir = "pwd"
            self.pwsd_dir = os.path.join ( os.path.dirname( os.path.abspath( __file__ ) ) , self.pwsd_dir )
            shutil.rmtree ( self.pwsd_dir )
        except:
            traceback.print_exc()

    def pass_exist_in_baidu ( self ):
        try:
            text_Found = False
            file_name_exists = False

            remote_path = "/pwd/"
            file_name = "t.json"
            print (" checking pass word exist  ", file_name, "\t remote path ", remote_path )
            command = [ "bypy", "search", file_name, remote_path ]
            proc =  subprocess.Popen( command  , stdout=subprocess.PIPE )
            while True:
                line = proc.stdout.readline()
                if not line:
                    break

                # print ( "test:", line.rstrip() )
                value = str ( line, "utf-8") .strip()
                print("subprocces ", value)

                if "Found" in value :
                    text_Found = True
                if text_Found and file_name in value:
                    file_name_exists = True
            self.flus_stdio ()
            if file_name_exists:
                return True
        except:
            print(traceback.format_exc())
            return  None
  

    def download_ps_to_pwd ( self ):
        try:
            self.remove_pwd_folder ()
            time.sleep (2)
            self.create_pwd_folder ()
            time.sleep ( 1 )
            command = [ "bypy","download","/pwd/", self.pwsd_dir ]
            subprocess.call(  command , creationflags=subprocess.CREATE_NEW_CONSOLE)
            self.flus_stdio()

        except:
            traceback.print_exc()


    def uplaod_ps_to_baidu ( self ):
        try:
            print ( "uploading from ", self.pwsd_dir )
            command = [ "bypy", "upload", self.pwsd_dir, "/pwd/" ]
            # subprocess.call(  command , creationflags=subprocess.CREATE_NEW_CONSOLE)
            
            proc = subprocess.Popen(command, stdout=subprocess.PIPE)

            while True:
                line = proc.stdout.readline()
                if not line:
                    break

                # print ( "test:", line.rstrip() )
                value = str(line, "utf-8")
                print("subprocces ", value)
        
            
        except:
            traceback.print_exc ( )

    
    def create_pass_at_pwd ( self , value ):
        self.remove_pwd_folder ()
        time.sleep (2)
        self.create_pwd_folder ()
        time.sleep (2)

        json_value = { "password": value }
        json_object = json.dumps( json_value )

        # Writing to sample.json
        path = os.path.join ( self.pwsd_dir , "t.json" )
        with open(  path, "w") as outfile:
            outfile.write(json_object)
        time.sleep(2)


    def read_pass_from_pwd ( self ):

        path = os.path.join ( self.pwsd_dir , "t.json" )

        f = open( path )
        # returns JSON object as
        # a dictionary
        data = json.load(f)
        password = data["password"]

        f.close()

        return password


    def flus_stdio ( self ):
        try:
            sys.stdout.flush()
        except:
            traceback.print_exc()
#!SECTION - password

## STUB - Class for Progress bar 
class ProgressDialog(QDialog):
    def __init__(self, parent=None, title="Download Progress"):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.progress_bar = QProgressBar(self)
        self.progress_bar.setGeometry(10, 10, 300, 30)

### start the app
if __name__ == "__main__":
    try:

        Database.open_db_connection ()
        DatabaseRunning.open_connection ()

        # #  first check password exists in baidu or not
        #
        """
        #FIXME -  later just remove the comment 
        check_pass = CheckPassWord ()
        pass_exists = check_pass.pass_exist_in_baidu()

        if pass_exists :
            check_pass.download_ps_to_pwd()
            pass
        else:
            check_pass.create_pass_at_pwd ("123456")
            check_pass.uplaod_ps_to_baidu()

        _PASSWORD = check_pass.read_pass_from_pwd()
        """
        #FIXME -  remove this line
        _PASSWORD = user_given_pass=  123456
        

        # for QT ui initiate system
        app = QApplication(sys.argv)

        # ui for password , 
        # by default it will ask user for password
        """#FIXME - 
        pass_w = Password()
        pass_w.setWindowIcon(QIcon('icon/Logo.png'))
        # pass_w.hide()

        user_given_pass = pass_w.get_password()
        
        

        # NOTE match pass word
        for i in range(1,3):
            if user_given_pass == _PASSWORD:
                break
            else:
                user_given_pass = pass_w.get_password ()
        """
        if not user_given_pass == _PASSWORD:
            print ("password does not match ")
            pass

        else:
            # loading style file <<Although it can be done via QFile >>
            with open('style.qss', 'r') as style_file:
                style_str = style_file.read()
            app.setStyleSheet(style_str)

            window = MainWindow()
            
            window.close_database ()
            DB_MGM.download_db_file ()
            DB_MGM.download_db_running ()

            window.show()
            window.setWindowIcon(QIcon('icon/Logo.png'))
            app.setWindowIcon(QIcon('icon/Logo.png'))
            app.exec()
            # sys.exit()
    except:
        print("has problem")
        traceback.print_exc()

    # time.sleep(100)
    # sys.exit()
    print("closing the application")
    window.close_database ()
    DB_MGM.upload_db_file ()
    time.sleep(3)


