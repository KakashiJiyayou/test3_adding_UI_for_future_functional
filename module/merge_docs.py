import os
import re
import subprocess
import sys
import time
import traceback
from shutil import copy

import pypandoc
from PIL import Image
from bs4 import BeautifulSoup
from urllib.parse import unquote

## Replace this with the actual path to soffice executable on your system
from docx import Document
from docx.shared import Inches




_soffice_path = ""

## output dir
_output_dir = ""

## single html file name
_single_file_name = "merged_html.html"

##  take a list of file paths
## if all paths exists return true
## even one path does not exist it will return False
def _check_all_file_exists ( file_list  ):
    return all ( os.path.exists( path )  for path in file_list )



## lets make given list to html files
def _convert_to_html ( file_list ):
    """
    Convert a list of DOC and DOCX files to HTML using LibreOffice.

    The conversion requires the following global parameters to be set:

    - `soffice_path`: Path to the LibreOffice executable.
    - `output_dir`: Directory where the HTML files will be saved.

    :param file_list: List of file paths to convert.
    """
    for file_path in file_list:
        _, file_extension = os.path.splitext(file_path)
        lower_extension = file_extension.lower()

        if lower_extension in ( '.doc', '.docx'):
            try:
                subprocess.run(
                    [_soffice_path,
                     "--headless",
                     "--convert-to",
                     'html',
                     file_path,
                     "--outdir",
                     _output_dir],
                    check=True
                )
                print(f"Conversion of {file_path} to HTML successful.")

                # change the file extension
                file_base, file_extension = os.path.splitext(file_path)
                file_path = file_base + '.html'
            except Exception as e:
                print(f"Error converting {file_path} to HTML: {e}")


## replace image
def _replace_image ( html_content, output_folder ):
    """
      Replace image tags in HTML content with placeholders .

      :param html_content: HTML content containing image tags.
      :param output_folder: Directory where the images will be copied.
      :return: Modified HTML content with placeholders.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    img_tags = soup.find_all('img')
    img_count = 0
    for img_tag in img_tags:
        img_src = img_tag.get('src')

        # got null
        if not img_src:
            continue

        img_count += 1
        try :
            print(" image replace issue ",img_src)
            img_src = unquote(img_src)

        except:
            traceback.print_exc( )
            sys.exit()

        img_filename = os.path.basename(img_src)

        dest_path = os.path.join(output_folder, img_filename)
        img_tag.replace_with( f'IMAG_SRC_{img_count}@@@{dest_path}' )

    modified_html = str(soup)
    return modified_html


## make one html file
def _make_all_html_into_one ( file_list, output_dir ):
    """
       Combine HTML files and embed images into a single HTML file.

       This function takes a list of file paths and combines the content of HTML files
       and embeds images using <img> tags into a single HTML file.
        - `single_file_name`: Path for new saved file.
       :param file_list: List of file paths to be combined and embedded.
       :param output_dir: Directory where the output HTML file will be saved.
    """
    combine_html = ""

    for path in file_list:
        if path.lower().endswith ( (".docx", ".doc") ):

            path = os.path.splitext( path )[0] + ".html"
            if not os.path.exists( path ):
                continue

            with open ( path, "r", encoding="utf-8" ) as file:
                html_content = file.read()
                soup = BeautifulSoup ( html_content, "html.parser" )
                body_content = str ( soup.body )
                combine_html += body_content
        elif path.lower().endswith ( ( '.jpg', '.jpeg', '.png', '.gif',
                                       '.tif', '.tiff', '.bmp',
                                       '.webp', '.svg', '.ico', '.heic' ) ):
            image_tag = f'<img src="{path}" alt="{os.path.basename(path)}">'
            combine_html += image_tag

        # Wrap the combined content in a complete HTML structure
        full_html = (
                '<!DOCTYPE html><html><head><meta charset="utf-8"/></head><body>'
                + combine_html
                + "</body></html>"
        )

        # replace images with text
        full_html = _replace_image( full_html, output_dir )

        output_path = os.path.join ( output_dir, _single_file_name )
        # Write the combined content to the output file
        with open(output_path, "w", encoding="utf-8") as file:
            file.write( full_html )

##
def find_image_path(image_path):
    if os.path.exists(image_path):
        return image_path
    return None

## convert html to docx add images
def _convert_html_to_doc_and_add_image ( output_dir ):
    """
        Convert HTML with embedded image placeholders to DOCX with actual images.

        This function takes an HTML file that contains embedded image placeholders and
        converts it into a DOCX file with actual images. The images are resized, and
        the resulting DOCX file is saved as 'FINAL_DOCX.docx'.

        :param output_dir: Directory containing the input HTML file and where the DOCX will be saved.
    """

    base_filename, _ = os.path.splitext(_single_file_name)
    output_file     = base_filename + ".docx"
    input_file      = os.path.join ( output_dir, _single_file_name )
    output_file     = os.path.join ( output_dir, output_file )
    new_doc_file = os.path.join ( output_dir, "FINAL_DOCX.docx" )

    pypandoc.convert(
        source=input_file,
        format='html',
        to='docx',
        outputfile=output_file,
        extra_args=['-RTS'])

    if not os.path.exists( output_file ):
        return None

    document = Document(output_file)
    index = 0
    for paragraph in document.paragraphs:
        line = paragraph.text.strip()

        # Use regular expression to find all occurrences of IMAG_SRC pattern
        img_matches = re.findall(r'IMAG_SRC_\d+@@@(.*?)(?=IMAG_SRC_\d+|$)', line)

        for img_match in img_matches:

            for i in range(0, index + len (img_matches) + 2 ):
                paragraph.text = paragraph.text.replace(f'IMAG_SRC_{index + i}@@@', '')
                print (f'_IMAG_SRC_{index +1 }')
            paragraph.text = paragraph.text.replace(img_match, '')

            print("img_match:", img_match , "\treplace ",f'IMAG_SRC_{index +1 }'  )
            img_match = img_match.strip()
            real_image_path = find_image_path( img_match)

            if real_image_path:
                img = Image.open( real_image_path )
                img = img.convert('RGB')

                temp_image_path = os.path.join(output_dir, f'temp_image{index}.jpg')
                index += 1
                img.save(temp_image_path, format='JPEG')
                run = paragraph.add_run()
                picture = run.add_picture(temp_image_path)

                # Set the maximum width and height (adjust as needed)
                max_width = Inches(8)
                max_height = Inches(6)


                # Calculate the aspect ratio
                aspect_ratio = picture.width / picture.height

                # Adjust width and height while preserving the aspect ratio
                if picture.width > max_width:
                    picture.width = max_width
                    picture.height = int(max_width / aspect_ratio)

                if picture.height > max_height:
                    picture.height = max_height
                    picture.width = int(max_height * aspect_ratio)

    document.save(new_doc_file)



## Take file list
#  Do many actions
def Merge ( file_list , output_dir, file_name, output_dir_generated_docx  ):
    """
        Merge and convert files from file_list to a single .docx file.
        
        Args:
            file_list (list): List of input file paths.
            output_dir (str): Directory where intermediate files will be generated.
            file_name (str): Name of the output .docx file.
            output_dir_generated_docx (str): Directory where the final .docx file will be copied.
    """
    ## first check file_list is none
    if file_list is None:
        print ( "can not merge None list ")
        return None

    ##
    if not _check_all_file_exists( file_list ):
        print ( "files missing in your computer.... " )
        return  None

    ##
    _convert_to_html( file_list )
    print( "Now check till now how the path looks like \n", file_list )

    ##
    _make_all_html_into_one ( file_list, output_dir )

    ##
    _convert_html_to_doc_and_add_image( output_dir )

    ##
    output_file     = "FINAL_DOCX.docx"
    source_path = os.path.join ( output_dir, output_file ) 
    result = os.path.exists ( source_path)

    if result:
        print ( "All process done" )
        destination_path = os.path.join(  output_dir_generated_docx, file_name)
        copy ( source_path, destination_path  )
    else:
        print ( "Genearating failled ")

    
    
    
    # copy file 







